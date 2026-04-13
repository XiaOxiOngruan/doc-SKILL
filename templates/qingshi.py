"""
请示模板
适用：向上级机关请求批准事项
"""

import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from utils.docx_utils import (
    load_format, apply_page_settings,
    add_heading1, add_heading2, add_body,
    add_separator, add_page_number
)


def build(
    title: str,
    doc_number: str,
    issuer: str,
    issue_date: str,
    recipient: str,
    background: str,
    request_items: list,
    reason: str,
    closing: str = "当否，请批示。",
    output_path: str = "请示.docx",
    format_path: str = None
):
    """
    生成请示类公文。

    参数：
        title          标题，如"关于申请XXX经费的请示"
        doc_number     发文字号
        issuer         发文机关
        issue_date     成文日期
        recipient      主送机关（请示只写一个主送机关）
        background     请示背景/事由段落
        request_items  请示事项列表，每项为字符串
        reason         请示理由
        closing        结语，默认"当否，请批示。"
        output_path    输出路径
        format_path    自定义格式路径
    """
    fmt = load_format(format_path)
    doc = Document()
    apply_page_settings(doc, fmt)

    # ── 版头 ────────────────────────────────────
    p_issuer = doc.add_paragraph()
    p_issuer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_issuer.add_run(issuer)
    run.font.name = "方正小标宋简体"
    run.font.size = Pt(22)
    from docx.shared import RGBColor
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "方正小标宋简体")
    rPr.insert(0, rFonts)

    add_separator(doc, fmt)

    # ── 字号、标题、主送 ────────────────────────
    add_body(doc, doc_number, fmt, indent=False)
    add_heading1(doc, title, fmt)
    add_body(doc, recipient + "：", fmt, indent=False)

    # ── 正文：背景 ──────────────────────────────
    add_heading2(doc, "一、基本情况", fmt)
    add_body(doc, background, fmt)

    # ── 请示事项 ────────────────────────────────
    add_heading2(doc, "二、请示事项", fmt)
    for i, item in enumerate(request_items, 1):
        label = ["（一）", "（二）", "（三）", "（四）", "（五）"]
        prefix = label[i - 1] if i <= len(label) else f"（{i}）"
        add_body(doc, prefix + item, fmt)

    # ── 请示理由 ────────────────────────────────
    add_heading2(doc, "三、请示理由", fmt)
    add_body(doc, reason, fmt)

    # ── 结语 ────────────────────────────────────
    add_body(doc, closing, fmt)

    # ── 落款 ────────────────────────────────────
    doc.add_paragraph()
    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_sign = p_sign.add_run(issuer + "\n" + issue_date)
    run_sign.font.name = "仿宋_GB2312"
    run_sign.font.size = Pt(16)
    rPr2 = run_sign._r.get_or_add_rPr()
    rFonts2 = OxmlElement("w:rFonts")
    rFonts2.set(qn("w:eastAsia"), "仿宋_GB2312")
    rPr2.insert(0, rFonts2)

    add_page_number(doc, fmt)

    doc.save(output_path)
    print(f"✓ 已生成：{os.path.abspath(output_path)}")
    return os.path.abspath(output_path)


if __name__ == "__main__":
    build(
        title="关于申请2024年度办公设备更新经费的请示",
        doc_number="办财〔2024〕5号",
        issuer="某某机关办公室",
        issue_date="2024年3月1日",
        recipient="某某机关",
        background="我单位现有办公设备购置于2015年，已超过使用年限，部分设备出现故障频繁、性能严重下降等问题，严重影响日常工作效率。",
        request_items=[
            "更新台式电脑20台，预计费用约20万元。",
            "更新打印机5台，预计费用约5万元。",
            "更新会议室投影设备2套，预计费用约3万元。",
        ],
        reason="以上设备更新工作对保障我单位正常运转具有重要意义，所需经费已纳入年度预算申报，请予审批。",
        output_path="examples/示例_请示.docx"
    )
