"""
lunwen.py  —  学术论文模板
支持：毕业论文、学位论文、期刊论文初稿、研究报告

结构：
  封面 → 中文摘要 → 英文摘要 → 目录（占位）→ 正文各章 → 参考文献 → 致谢

用法：
    python3 templates/lunwen.py
"""

import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from utils.docx_utils import (
    load_format, apply_page_settings, add_page_number
)

# ─────────────────────────────────────────────
# 内部工具函数
# ─────────────────────────────────────────────

def _set_font(run, font_name: str, size_pt: float, bold: bool = False,
              color_rgb: tuple = None, italic: bool = False):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    run.font.name = font_name
    if color_rgb:
        run.font.color.rgb = RGBColor(*color_rgb)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"),    font_name)
    rFonts.set(qn("w:hAnsi"),    font_name)
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


def _set_spacing(para, line_pt: float, space_before: float = 0,
                 space_after: float = 0, indent_chars: int = 0,
                 font_size_pt: float = 12):
    pf = para.paragraph_format
    pf.line_spacing      = Pt(line_pt)
    pf.line_spacing_rule = 4   # EXACTLY
    pf.space_before      = Pt(space_before)
    pf.space_after       = Pt(space_after)
    if indent_chars:
        pf.first_line_indent = Pt(font_size_pt * indent_chars)


def _add_text(doc, text: str, font_name: str, size_pt: float,
              bold=False, italic=False, align="left",
              line_pt=22, indent=0, space_after=6, color_rgb=None):
    p = doc.add_paragraph()
    align_map = {
        "center":  WD_ALIGN_PARAGRAPH.CENTER,
        "left":    WD_ALIGN_PARAGRAPH.LEFT,
        "right":   WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    p.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    _set_spacing(p, line_pt, space_after=space_after,
                 indent_chars=indent, font_size_pt=size_pt)
    if text:
        run = p.add_run(text)
        _set_font(run, font_name, size_pt, bold=bold,
                  italic=italic, color_rgb=color_rgb)
    return p


def _page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(__import__("docx.enum.text", fromlist=["WD_BREAK"]).WD_BREAK.PAGE)
    return p


# ─────────────────────────────────────────────
# 各结构块
# ─────────────────────────────────────────────

def add_cover(doc, fmt, title: str, author: str, institution: str,
              supervisor: str = "", degree: str = "", date: str = "",
              subtitle: str = ""):
    """封面页"""
    f_body = fmt["fonts"]["body"]
    f_h1   = fmt["fonts"]["heading1"]
    ls     = fmt["paragraph"]["line_spacing_pt"]

    # 顶部机构名
    _add_text(doc, institution, f_body["name"], f_body["size_pt"],
              align="center", line_pt=ls, space_after=0)

    # 大标题
    doc.add_paragraph()
    _add_text(doc, title, fmt["fonts"].get("cover_title", f_h1)["name"],
              fmt["fonts"].get("cover_title", f_h1)["size_pt"],
              bold=True, align="center", line_pt=ls + 4, space_after=4)

    # 副标题
    if subtitle:
        _add_text(doc, f"——{subtitle}", f_body["name"], f_body["size_pt"] + 1,
                  align="center", line_pt=ls, space_after=8)

    doc.add_paragraph()

    # 信息栏
    info_items = [("作者", author)]
    if supervisor: info_items.append(("指导教师", supervisor))
    if degree:     info_items.append(("申请学位", degree))
    if date:       info_items.append(("完成时间", date))

    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_spacing(p, ls, space_after=4)
        run_label = p.add_run(f"{label}：")
        _set_font(run_label, f_body["name"], f_body["size_pt"])
        run_value = p.add_run(value)
        _set_font(run_value, f_body["name"], f_body["size_pt"], bold=True)

    _page_break(doc)


def add_abstract_zh(doc, fmt, abstract_text: str, keywords: list):
    """中文摘要页"""
    f_body = fmt["fonts"]["abstract"]
    f_h1   = fmt["fonts"]["heading1"]
    ls     = fmt["paragraph"]["line_spacing_pt"]
    indent = fmt["paragraph"]["first_line_indent_chars"]

    _add_text(doc, "摘  要", f_h1["name"], f_h1["size_pt"],
              bold=True, align="center", line_pt=ls, space_after=12)

    _add_text(doc, abstract_text, f_body["name"], f_body["size_pt"],
              align="justify", line_pt=ls, indent=indent,
              space_after=fmt["paragraph"]["space_after_pt"])

    doc.add_paragraph()
    p = doc.add_paragraph()
    _set_spacing(p, ls, indent_chars=0, font_size_pt=f_body["size_pt"])
    run_label = p.add_run("关键词：")
    _set_font(run_label, f_body["name"], f_body["size_pt"], bold=True)
    run_kw = p.add_run("；".join(keywords))
    _set_font(run_kw, f_body["name"], f_body["size_pt"])

    _page_break(doc)


def add_abstract_en(doc, fmt, title_en: str, abstract_text_en: str,
                    keywords_en: list):
    """英文摘要页"""
    f_en   = fmt["fonts"].get("abstract_en", {"name": "Times New Roman", "size_pt": 12})
    f_h1   = fmt["fonts"]["heading1"]
    ls     = fmt["paragraph"]["line_spacing_pt"]
    indent = fmt["paragraph"]["first_line_indent_chars"]

    # 英文标题
    _add_text(doc, title_en, f_h1["name"], f_h1["size_pt"],
              bold=True, align="center", line_pt=ls, space_after=4)

    _add_text(doc, "Abstract", f_en["name"], f_en["size_pt"],
              bold=True, align="center", line_pt=ls, space_after=12)

    _add_text(doc, abstract_text_en, f_en["name"], f_en["size_pt"],
              align="justify", line_pt=ls, indent=indent)

    doc.add_paragraph()
    p = doc.add_paragraph()
    _set_spacing(p, ls, font_size_pt=f_en["size_pt"])
    run_label = p.add_run("Keywords: ")
    _set_font(run_label, f_en["name"], f_en["size_pt"], bold=True)
    run_kw = p.add_run("; ".join(keywords_en))
    _set_font(run_kw, f_en["name"], f_en["size_pt"])

    _page_break(doc)


def add_toc_placeholder(doc, fmt):
    """目录占位页（实际目录在 Word 中右键更新域生成）"""
    f_h1 = fmt["fonts"]["heading1"]
    f_b  = fmt["fonts"]["body"]
    ls   = fmt["paragraph"]["line_spacing_pt"]

    _add_text(doc, "目  录", f_h1["name"], f_h1["size_pt"],
              bold=True, align="center", line_pt=ls, space_after=12)

    p = doc.add_paragraph()
    _set_spacing(p, ls, font_size_pt=f_b["size_pt"])
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn
    # 插入 TOC 域代码
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(_qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(_qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(_qn("w:fldCharType"), "separate")
    placeholder_r = OxmlElement("w:r")
    placeholder_t = OxmlElement("w:t")
    placeholder_t.text = "【在 Word 中右键此处 → 更新域，即可生成目录】"
    placeholder_r.append(placeholder_t)
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(_qn("w:fldCharType"), "end")
    r = OxmlElement("w:r")
    r.append(fldChar1)
    r_instr = OxmlElement("w:r")
    r_instr.append(instr)
    r_sep = OxmlElement("w:r")
    r_sep.append(fldChar2)
    r_end = OxmlElement("w:r")
    r_end.append(fldChar3)
    p._p.append(r)
    p._p.append(r_instr)
    p._p.append(r_sep)
    p._p.append(placeholder_r)
    p._p.append(r_end)

    _page_break(doc)


def add_chapter(doc, fmt, chapter_title: str, sections: list,
                chapter_number: int = 1):
    """
    正文章节。

    sections 格式：
    [
      {
        "title": "1.1 研究背景",      # 节标题，空字符串则不加节标题
        "body":  ["段落1", "段落2"],   # 正文段落列表
        "sub": [                       # 可选子节
          {"title": "1.1.1 ...", "body": ["..."]}
        ]
      }
    ]
    """
    f_h1   = fmt["fonts"]["heading1"]
    f_h2   = fmt["fonts"]["heading2"]
    f_h3   = fmt["fonts"]["heading3"]
    f_body = fmt["fonts"]["body"]
    ls     = fmt["paragraph"]["line_spacing_pt"]
    indent = fmt["paragraph"]["first_line_indent_chars"]
    sa     = fmt["paragraph"]["space_after_pt"]

    # 章标题
    _add_text(doc, chapter_title, f_h1["name"], f_h1["size_pt"],
              bold=f_h1.get("bold", True), align="center",
              line_pt=ls, space_after=12)

    for sec in sections:
        # 节标题
        if sec.get("title"):
            _add_text(doc, sec["title"], f_h2["name"], f_h2["size_pt"],
                      bold=f_h2.get("bold", True), align="left",
                      line_pt=ls, space_after=6)

        for para in sec.get("body", []):
            _add_text(doc, para, f_body["name"], f_body["size_pt"],
                      align="justify", line_pt=ls,
                      indent=indent, space_after=sa)

        # 子节
        for sub in sec.get("sub", []):
            if sub.get("title"):
                _add_text(doc, sub["title"], f_h3["name"], f_h3["size_pt"],
                          bold=f_h3.get("bold", True), align="left",
                          line_pt=ls, space_after=4)
            for para in sub.get("body", []):
                _add_text(doc, para, f_body["name"], f_body["size_pt"],
                          align="justify", line_pt=ls,
                          indent=indent, space_after=sa)


def add_references(doc, fmt, references: list):
    """
    参考文献列表。

    references 为字符串列表，每项为一条已格式化的参考文献，如：
    "[1] 作者. 书名[M]. 出版社, 年份."
    """
    f_h1  = fmt["fonts"]["heading1"]
    f_ref = fmt["fonts"].get("reference", fmt["fonts"]["body"])
    ls    = fmt["paragraph"]["line_spacing_pt"]

    _page_break(doc)
    _add_text(doc, "参考文献", f_h1["name"], f_h1["size_pt"],
              bold=True, align="center", line_pt=ls, space_after=12)

    for ref in references:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing      = Pt(ls)
        pf.line_spacing_rule = 4
        pf.space_after       = Pt(3)
        pf.first_line_indent = Pt(0)
        pf.left_indent       = Pt(f_ref["size_pt"] * 2)  # 悬挂缩进
        pf.first_line_indent = Pt(-f_ref["size_pt"] * 2)
        run = p.add_run(ref)
        _set_font(run, f_ref["name"], f_ref["size_pt"])


def add_acknowledgement(doc, fmt, text: str):
    """致谢"""
    f_h1   = fmt["fonts"]["heading1"]
    f_body = fmt["fonts"]["body"]
    ls     = fmt["paragraph"]["line_spacing_pt"]
    indent = fmt["paragraph"]["first_line_indent_chars"]

    _page_break(doc)
    _add_text(doc, "致  谢", f_h1["name"], f_h1["size_pt"],
              bold=True, align="center", line_pt=ls, space_after=12)

    for para in text.split("\n"):
        if para.strip():
            _add_text(doc, para.strip(), f_body["name"], f_body["size_pt"],
                      align="justify", line_pt=ls, indent=indent, space_after=6)


# ─────────────────────────────────────────────
# 主构建函数
# ─────────────────────────────────────────────

def build(
    title: str,
    author: str,
    institution: str = "",
    supervisor: str = "",
    degree: str = "",
    date: str = "",
    subtitle: str = "",
    abstract_zh: str = "",
    keywords_zh: list = None,
    title_en: str = "",
    abstract_en: str = "",
    keywords_en: list = None,
    chapters: list = None,
    references: list = None,
    acknowledgement: str = "",
    include_toc: bool = True,
    output_path: str = "论文.docx",
    format_path: str = None
):
    """
    生成学术论文 .docx 文件。

    参数：
        title           论文标题（中文）
        author          作者姓名
        institution     所在机构/学校
        supervisor      指导教师（可选）
        degree          申请学位，如"工学硕士"（可选）
        date            完成时间，如"2024年6月"（可选）
        subtitle        副标题（可选）
        abstract_zh     中文摘要正文
        keywords_zh     中文关键词列表，如 ["深度学习", "目标检测"]
        title_en        英文标题（可选，不传则跳过英文摘要页）
        abstract_en     英文摘要
        keywords_en     英文关键词列表
        chapters        章节列表，格式见 add_chapter 说明
        references      参考文献列表，每项为已格式化字符串
        acknowledgement 致谢正文（空则跳过）
        include_toc     是否插入目录占位（默认 True）
        output_path     输出文件路径
        format_path     自定义格式文件路径（.json/.yaml），默认使用论文格式
    """
    if format_path is None:
        default_thesis_fmt = os.path.join(
            os.path.dirname(__file__), "..", "formats", "academic_thesis.json"
        )
        if os.path.exists(default_thesis_fmt):
            fmt = load_format(default_thesis_fmt)
        else:
            fmt = load_format()
    else:
        fmt = load_format(format_path)

    doc = Document()
    apply_page_settings(doc, fmt)

    # ── 页眉：论文标题居中 ────────────────────
    section = doc.sections[0]
    header  = section.header
    if header.paragraphs:
        ph = header.paragraphs[0]
    else:
        ph = header.add_paragraph()
    ph.clear()
    ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_h = ph.add_run(title)
    _set_font(run_h, fmt["fonts"]["header"]["name"],
              fmt["fonts"]["header"]["size_pt"])

    # ── 封面 ─────────────────────────────────
    add_cover(doc, fmt, title, author, institution,
              supervisor=supervisor, degree=degree,
              date=date, subtitle=subtitle)

    # ── 中文摘要 ──────────────────────────────
    if abstract_zh:
        add_abstract_zh(doc, fmt, abstract_zh, keywords_zh or [])

    # ── 英文摘要 ──────────────────────────────
    if title_en and abstract_en:
        add_abstract_en(doc, fmt, title_en, abstract_en, keywords_en or [])

    # ── 目录 ─────────────────────────────────
    if include_toc:
        add_toc_placeholder(doc, fmt)

    # ── 正文各章 ──────────────────────────────
    for i, chapter in enumerate(chapters or [], start=1):
        if i > 1:
            _page_break(doc)
        add_chapter(doc, fmt,
                    chapter_title=chapter["title"],
                    sections=chapter.get("sections", []),
                    chapter_number=i)

    # ── 参考文献 ──────────────────────────────
    if references:
        add_references(doc, fmt, references)

    # ── 致谢 ─────────────────────────────────
    if acknowledgement:
        add_acknowledgement(doc, fmt, acknowledgement)

    # ── 页码 ─────────────────────────────────
    add_page_number(doc, fmt)

    doc.save(output_path)
    print(f"✓ 已生成：{os.path.abspath(output_path)}")
    return os.path.abspath(output_path)


# ─────────────────────────────────────────────
# 示例
# ─────────────────────────────────────────────

if __name__ == "__main__":
    build(
        title="基于深度学习的中文文本分类研究",
        author="张三",
        institution="某某大学计算机学院",
        supervisor="李四 教授",
        degree="工学硕士",
        date="2024年6月",
        abstract_zh=(
            "随着自然语言处理技术的快速发展，深度学习方法在中文文本分类任务中取得了显著成效。"
            "本文提出了一种基于预训练语言模型的中文文本分类框架，通过引入多粒度注意力机制，"
            "有效捕捉文本的局部和全局语义信息。在多个公开数据集上的实验表明，"
            "所提方法在准确率和F1值等指标上均优于现有基线方法，具有较好的实用价值。"
        ),
        keywords_zh=["深度学习", "文本分类", "预训练模型", "注意力机制", "自然语言处理"],
        title_en="Research on Chinese Text Classification Based on Deep Learning",
        abstract_en=(
            "With the rapid development of natural language processing, deep learning methods "
            "have achieved remarkable results in Chinese text classification tasks. "
            "This paper proposes a Chinese text classification framework based on pre-trained "
            "language models. By introducing a multi-granularity attention mechanism, "
            "the framework effectively captures both local and global semantic information. "
            "Experiments on multiple public datasets demonstrate that the proposed method "
            "outperforms existing baselines on accuracy and F1 metrics."
        ),
        keywords_en=["Deep Learning", "Text Classification", "Pre-trained Model",
                     "Attention Mechanism", "NLP"],
        chapters=[
            {
                "title": "第一章  绪论",
                "sections": [
                    {
                        "title": "1.1 研究背景与意义",
                        "body": [
                            "近年来，互联网的快速发展产生了海量文本数据，如何自动、准确地对这些文本进行分类成为信息处理领域的重要研究课题。",
                            "文本分类是自然语言处理的基础任务之一，在新闻分类、情感分析、垃圾邮件过滤等领域具有广泛的应用前景。",
                        ],
                        "sub": [
                            {
                                "title": "1.1.1 研究背景",
                                "body": ["中文文本具有词边界模糊、一词多义等特点，给文本分类带来了额外挑战。"]
                            }
                        ]
                    },
                    {
                        "title": "1.2 研究现状",
                        "body": [
                            "传统文本分类方法主要依赖人工特征工程，包括词袋模型、TF-IDF等方法。随着深度学习的兴起，CNN、LSTM等神经网络模型逐渐成为主流。",
                            "近年来，以BERT为代表的预训练语言模型的出现进一步提升了文本分类的性能上限。",
                        ]
                    },
                    {
                        "title": "1.3 本文组织结构",
                        "body": [
                            "本文共分为五章。第一章介绍研究背景与意义；第二章回顾相关工作；第三章介绍所提方法；第四章进行实验验证；第五章总结全文并展望未来工作。"
                        ]
                    }
                ]
            },
            {
                "title": "第二章  相关工作",
                "sections": [
                    {
                        "title": "2.1 传统文本分类方法",
                        "body": [
                            "朴素贝叶斯、支持向量机等传统机器学习方法在文本分类中得到广泛应用，但这些方法依赖人工构建特征，泛化能力有限。"
                        ]
                    },
                    {
                        "title": "2.2 深度学习方法",
                        "body": [
                            "TextCNN通过卷积操作提取局部特征，在短文本分类上表现出色。BiLSTM利用双向循环结构捕捉上下文信息，适合处理长文本。"
                        ]
                    }
                ]
            },
            {
                "title": "第三章  方法",
                "sections": [
                    {
                        "title": "3.1 模型总体框架",
                        "body": [
                            "本文提出的模型以BERT为骨干网络，在其输出层引入多粒度注意力模块，分别从字级、词级和句级三个粒度提取文本特征。"
                        ]
                    },
                    {
                        "title": "3.2 多粒度注意力机制",
                        "body": [
                            "给定输入序列，模型首先通过BERT编码器得到上下文表示，随后分别在不同粒度上计算注意力权重，加权融合后输入分类层。"
                        ]
                    }
                ]
            },
            {
                "title": "第四章  实验",
                "sections": [
                    {
                        "title": "4.1 实验设置",
                        "body": [
                            "实验在THUCNews、Fudan等中文文本分类数据集上进行，采用准确率（Accuracy）、宏平均F1（Macro-F1）作为评价指标。"
                        ]
                    },
                    {
                        "title": "4.2 实验结果",
                        "body": [
                            "在所有数据集上，本文方法均优于基线方法。在THUCNews数据集上，准确率达到97.3%，较最优基线提升0.8个百分点。"
                        ]
                    }
                ]
            },
            {
                "title": "第五章  总结与展望",
                "sections": [
                    {
                        "title": "5.1 工作总结",
                        "body": [
                            "本文针对中文文本分类问题，提出了基于多粒度注意力的预训练模型微调方法，在多个数据集上验证了方法的有效性。"
                        ]
                    },
                    {
                        "title": "5.2 未来工作",
                        "body": [
                            "后续工作将探索更轻量化的模型结构，以降低推理成本，并尝试将方法迁移至跨语言文本分类场景。"
                        ]
                    }
                ]
            }
        ],
        references=[
            "[1] Devlin J, Chang M W, Lee K, et al. BERT: Pre-training of deep bidirectional transformers for language understanding[C]. NAACL, 2019: 4171-4186.",
            "[2] Kim Y. Convolutional neural networks for sentence classification[C]. EMNLP, 2014: 1746-1751.",
            "[3] 李航. 统计学习方法[M]. 第2版. 北京: 清华大学出版社, 2019.",
            "[4] Liu Y, Ott M, Goyal N, et al. RoBERTa: A robustly optimized BERT pretraining approach[J]. arXiv preprint arXiv:1907.11692, 2019.",
        ],
        acknowledgement=(
            "首先，衷心感谢我的导师李四教授在本论文研究过程中给予的悉心指导和大力支持。\n"
            "感谢实验室各位同学在学习和生活中的帮助与陪伴。\n"
            "最后，向家人表达最深切的感谢，感谢他们一直以来的理解和鼓励。"
        ),
        output_path="examples/示例_论文.docx"
    )
