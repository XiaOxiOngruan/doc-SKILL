"""
通知模板
适用：转发文件、部署工作、传达事项等
"""

import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx import Document
from utils.docx_utils import (
    load_format, apply_page_settings,
    add_heading1, add_heading2, add_body,
    add_separator, add_page_number
)


def build(
    title: str,
    doc_number: str,
    issuer: str,
    issue_date: str,
    recipients: str,
    sections: list,
    closing: str = "",
    output_path: str = "通知.docx",
    format_path: str = None
):
    """
    生成通知类公文。

    参数：
        title        文件标题，如"关于开展XXX工作的通知"
        doc_number   发文字号，如"国办发〔2024〕1号"
        issuer       发文机关，如"国务院办公厅"
        issue_date   成文日期，如"2024年1月1日"
        recipients   主送机关，如"各省、自治区、直辖市人民政府"
        sections     正文章节列表，格式：
                     [
                       {"heading": "一、工作背景", "body": ["段落1", "段落2"]},
                       {"heading": "二、主要任务", "body": ["段落1"]},
                     ]
        closing      附注或结语，如"请遵照执行。"
        output_path  输出文件路径
        format_path  自定义格式文件路径（.json/.yaml），不传则用 GB/T 9704 默认格式
    """
    fmt = load_format(format_path)
    doc = Document()
    apply_page_settings(doc, fmt)

    # ── 版头：发文机关 ──────────────────────────
    p_issuer = doc.add_paragraph()
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p_issuer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_issuer.add_run(issuer)
    run.font.name = "方正小标宋简体"
    run.font.size = Pt(22)
    run.font.color.rgb = __import__("docx.shared", fromlist=["RGBColor"]).RGBColor(0xCC, 0x00, 0x00)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "方正小标宋简体")
    rPr.insert(0, rFonts)

    # ── 红色分隔线 ──────────────────────────────
    add_separator(doc, fmt)

    # ── 发文字号 ────────────────────────────────
    add_body(doc, doc_number, fmt, indent=False)

    # ── 文件标题 ────────────────────────────────
    add_heading1(doc, title, fmt)

    # ── 主送机关 ────────────────────────────────
    add_body(doc, recipients + "：", fmt, indent=False)

    # ── 正文各节 ────────────────────────────────
    for section in sections:
        if section.get("heading"):
            add_heading2(doc, section["heading"], fmt)
        for para in section.get("body", []):
            add_body(doc, para, fmt)

    # ── 结语 ────────────────────────────────────
    if closing:
        add_body(doc, closing, fmt)

    # ── 落款：发文机关 + 日期 ───────────────────
    doc.add_paragraph()
    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_sign = p_sign.add_run(issuer + "\n" + issue_date)
    run_sign.font.name = "仿宋_GB2312"
    run_sign.font.size = Pt(16)
    rPr2 = run_sign._r.get_or_add_rPr()
    rFonts2 = OxmlElement("w:rFonts")
    rFonts2.set(qn("w:eastAsia"), "仿宋_GB2312")
    rPr2.insert(0, rFonts2)

    # ── 页码 ────────────────────────────────────
    add_page_number(doc, fmt)

    doc.save(output_path)
    print(f"✓ 已生成：{os.path.abspath(output_path)}")
    return os.path.abspath(output_path)


if __name__ == "__main__":
    build(
        title="关于做好2024年度安全生产工作的通知",
        doc_number="办发〔2024〕1号",
        issuer="某某机关办公室",
        issue_date="2024年1月15日",
        recipients="各部门、各单位",
        sections=[
            {
                "heading": "一、总体要求",
                "body": [
                    "深入贯彻习近平总书记关于安全生产的重要论述，坚持人民至上、生命至上，强化安全生产责任制，确保全年安全生产形势持续稳定好转。",
                ]
            },
            {
                "heading": "二、重点任务",
                "body": [
                    "（一）完善安全生产责任体系。各单位主要负责同志要认真履行安全生产第一责任人职责，层层压实责任。",
                    "（二）开展安全生产专项整治。聚焦重点行业领域，深入排查整治各类安全隐患，确保整改到位。",
                    "（三）加强安全生产宣传教育。定期组织安全培训，提高全员安全意识和应急处置能力。",
                ]
            },
            {
                "heading": "三、工作要求",
                "body": [
                    "各部门要高度重视，认真组织实施，于每季度末将工作落实情况报办公室汇总。",
                ]
            }
        ],
        closing="请遵照执行。",
        output_path="examples/示例_通知.docx"
    )
