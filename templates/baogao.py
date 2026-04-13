"""
报告模板
适用：工作报告、情况报告、调研报告等
"""

import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from utils.docx_utils import (
    load_format, apply_page_settings,
    add_heading1, add_heading2, add_heading3, add_body,
    add_separator, add_page_number, add_image
)


def build(
    title: str,
    doc_number: str,
    issuer: str,
    issue_date: str,
    recipient: str,
    sections: list,
    attachments: list = None,
    images: list = None,
    output_path: str = "报告.docx",
    format_path: str = None
):
    """
    生成报告类公文。

    参数：
        title        标题
        doc_number   发文字号
        issuer       发文机关
        issue_date   成文日期
        recipient    主送机关
        sections     章节列表，支持三级嵌套：
                     [
                       {
                         "heading": "一、工作情况",
                         "body": ["段落1"],
                         "sub": [
                           {"heading": "（一）具体工作", "body": ["段落"]},
                         ]
                       }
                     ]
        attachments  附件列表，如 ["1.XXX材料", "2.XXX表格"]
        images       图片列表，如 [{"path": "xxx.png", "caption": "图1 流程图"}]
        output_path  输出路径
        format_path  自定义格式路径
    """
    fmt = load_format(format_path)
    doc = Document()
    apply_page_settings(doc, fmt)

    # ── 版头 ────────────────────────────────────
    p_issuer = doc.add_paragraph()
    p_issuer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_issuer.add_run(issuer)
    run.font.name = "方正小标宋简体"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "方正小标宋简体")
    rPr.insert(0, rFonts)

    add_separator(doc, fmt)
    add_body(doc, doc_number, fmt, indent=False)
    add_heading1(doc, title, fmt)
    add_body(doc, recipient + "：", fmt, indent=False)

    # ── 正文 ────────────────────────────────────
    for section in sections:
        if section.get("heading"):
            add_heading2(doc, section["heading"], fmt)
        for para in section.get("body", []):
            add_body(doc, para, fmt)

        # 三级子章节
        for sub in section.get("sub", []):
            if sub.get("heading"):
                add_heading3(doc, sub["heading"], fmt)
            for para in sub.get("body", []):
                add_body(doc, para, fmt)

    # ── 图片 ────────────────────────────────────
    if images:
        for img in images:
            if os.path.exists(img["path"]):
                add_image(doc, img["path"])
                if img.get("caption"):
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_cap = p_cap.add_run(img["caption"])
                    run_cap.font.name = "仿宋_GB2312"
                    run_cap.font.size = Pt(14)

    # ── 附件说明 ────────────────────────────────
    if attachments:
        add_body(doc, "附件：", fmt, indent=False)
        for att in attachments:
            add_body(doc, att, fmt)

    # ── 落款 ────────────────────────────────────
    doc.add_paragraph()
    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_sign = p_sign.add_run(issuer + "\n" + issue_date)
    run_sign.font.name = "仿宋_GB2312"
    run_sign.font.size = Pt(16)
    rPr2 = run_sign._r.get_or_add_rPr()
    rFonts2 = OxmlElement("w:rFonts")
    rFonts2.set(qn("w:eastAsia"), "仿宋_GB2312")
    rPr2.insert(0, rFonts2)

    add_page_number(doc, fmt)

    doc.save(output_path)
    print(f"✓ 已生成：{os.path.abspath(output_path)}")
    return os.path.abspath(output_path)


if __name__ == "__main__":
    build(
        title="关于2024年上半年工作情况的报告",
        doc_number="办报〔2024〕3号",
        issuer="某某机关办公室",
        issue_date="2024年7月5日",
        recipient="某某机关",
        sections=[
            {
                "heading": "一、上半年工作情况",
                "body": ["上半年，我单位认真贯彻落实上级部署要求，扎实推进各项工作，总体运行平稳、成效明显。"],
                "sub": [
                    {
                        "heading": "（一）深化改革工作扎实推进",
                        "body": [
                            "认真落实改革任务清单，完成改革事项12项，其中重点改革5项均按时序推进。"
                        ]
                    },
                    {
                        "heading": "（二）业务工作稳步提升",
                        "body": [
                            "累计完成业务量XXX，同比增长XX%，创近年来同期最好水平。"
                        ]
                    }
                ]
            },
            {
                "heading": "二、存在的问题",
                "body": [
                    "一是部分工作推进不够平衡，个别指标完成进度滞后；二是基层能力建设仍需加强。"
                ],
                "sub": []
            },
            {
                "heading": "三、下半年工作安排",
                "body": [
                    "下半年，将重点做好以下工作：一是加快推进重点改革任务落地；二是全力完成年度目标任务。"
                ],
                "sub": []
            }
        ],
        attachments=["1.2024年上半年重点工作完成情况表"],
        output_path="examples/示例_报告.docx"
    )
