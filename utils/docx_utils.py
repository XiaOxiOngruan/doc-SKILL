"""
公文排版工具函数
支持通用国标 GB/T 9704 及自定义格式
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
import os

try:
    import yaml
    _YAML_AVAILABLE = True
except ImportError:
    _YAML_AVAILABLE = False


# ─────────────────────────────────────────────
# 默认格式：GB/T 9704 国家标准公文格式
# ─────────────────────────────────────────────
DEFAULT_FORMAT = {
    "name": "GB/T 9704 国家标准",
    "page": {
        "top_cm": 3.7,
        "bottom_cm": 3.5,
        "left_cm": 2.8,
        "right_cm": 2.6,
        "page_size": "A4"
    },
    "fonts": {
        "heading1": {"name": "方正小标宋简体", "size_pt": 22, "bold": False, "align": "center"},
        "heading2": {"name": "黑体",           "size_pt": 16, "bold": False, "align": "left"},
        "heading3": {"name": "楷体_GB2312",    "size_pt": 16, "bold": False, "align": "left"},
        "body":     {"name": "仿宋_GB2312",    "size_pt": 16, "bold": False, "align": "justify"},
        "header":   {"name": "仿宋_GB2312",    "size_pt": 14, "bold": False},
        "footer":   {"name": "宋体",           "size_pt": 14, "bold": False}
    },
    "paragraph": {
        "line_spacing_pt": 28,
        "first_line_indent_chars": 2,
        "space_before_pt": 0,
        "space_after_pt": 0
    },
    "structure": {
        "levels": ["一、", "（一）", "1.", "（1）"]
    },
    "page_number": {
        "format": "—{n}—",
        "odd_align": "right",
        "even_align": "left"
    }
}


def load_format(format_path: str = None) -> dict:
    """
    加载格式配置。
    - 不传参数：使用内置 GB/T 9704 默认格式
    - 传入 .json 或 .yaml/.yml 路径：加载自定义格式，缺失字段自动回退到默认值
    """
    if not format_path:
        return DEFAULT_FORMAT.copy()

    ext = os.path.splitext(format_path)[1].lower()
    with open(format_path, "r", encoding="utf-8") as f:
        if ext == ".json":
            custom = json.load(f)
        elif ext in (".yaml", ".yml"):
            if not _YAML_AVAILABLE:
                raise ImportError("读取 YAML 格式需要安装 pyyaml：pip install pyyaml")
            custom = yaml.safe_load(f)
        else:
            raise ValueError(f"不支持的格式文件类型：{ext}，请使用 .json 或 .yaml")

    # 深度合并，自定义覆盖默认
    merged = _deep_merge(DEFAULT_FORMAT, custom)
    return merged


def _deep_merge(base: dict, override: dict) -> dict:
    result = base.copy()
    for k, v in override.items():
        if k in result and isinstance(result[k], dict) and isinstance(v, dict):
            result[k] = _deep_merge(result[k], v)
        else:
            result[k] = v
    return result


def apply_page_settings(doc: Document, fmt: dict):
    """设置页边距"""
    p = fmt["page"]
    section = doc.sections[0]
    section.top_margin    = Cm(p["top_cm"])
    section.bottom_margin = Cm(p["bottom_cm"])
    section.left_margin   = Cm(p["left_cm"])
    section.right_margin  = Cm(p["right_cm"])


def _set_run_font(run, font_name: str, size_pt: float, bold: bool = False, color_rgb: tuple = None):
    """设置 run 的字体（中英文统一）"""
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = font_name
    if color_rgb:
        run.font.color.rgb = RGBColor(*color_rgb)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


def _set_line_spacing(paragraph, spacing_pt: float, space_before: float = 0, space_after: float = 0):
    """设置固定行距"""
    pf = paragraph.paragraph_format
    pf.line_spacing = Pt(spacing_pt)
    pf.line_spacing_rule = 4   # WD_LINE_SPACING.EXACTLY
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)


def _align_map(align_str: str):
    m = {
        "center":  WD_ALIGN_PARAGRAPH.CENTER,
        "left":    WD_ALIGN_PARAGRAPH.LEFT,
        "right":   WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return m.get(align_str, WD_ALIGN_PARAGRAPH.LEFT)


# ─────────────────────────────────────────────
# 段落添加函数
# ─────────────────────────────────────────────

def add_heading1(doc: Document, text: str, fmt: dict) -> object:
    """一级标题（文件标题）"""
    cfg = fmt["fonts"]["heading1"]
    para = fmt["paragraph"]
    p = doc.add_paragraph()
    _set_line_spacing(p, para["line_spacing_pt"], para["space_before_pt"], para["space_after_pt"])
    p.alignment = _align_map(cfg["align"])
    run = p.add_run(text)
    _set_run_font(run, cfg["name"], cfg["size_pt"], cfg["bold"])
    return p


def add_heading2(doc: Document, text: str, fmt: dict) -> object:
    """二级标题"""
    cfg = fmt["fonts"]["heading2"]
    para = fmt["paragraph"]
    p = doc.add_paragraph()
    _set_line_spacing(p, para["line_spacing_pt"], para["space_before_pt"], para["space_after_pt"])
    p.alignment = _align_map(cfg["align"])
    run = p.add_run(text)
    _set_run_font(run, cfg["name"], cfg["size_pt"], cfg["bold"])
    return p


def add_heading3(doc: Document, text: str, fmt: dict) -> object:
    """三级标题"""
    cfg = fmt["fonts"]["heading3"]
    para = fmt["paragraph"]
    p = doc.add_paragraph()
    _set_line_spacing(p, para["line_spacing_pt"], para["space_before_pt"], para["space_after_pt"])
    p.alignment = _align_map(cfg["align"])
    run = p.add_run(text)
    _set_run_font(run, cfg["name"], cfg["size_pt"], cfg["bold"])
    return p


def add_body(doc: Document, text: str, fmt: dict, indent: bool = True) -> object:
    """正文段落，首行缩进 2 字符"""
    cfg = fmt["fonts"]["body"]
    para = fmt["paragraph"]
    p = doc.add_paragraph()
    _set_line_spacing(p, para["line_spacing_pt"], para["space_before_pt"], para["space_after_pt"])
    p.alignment = _align_map(cfg["align"])
    if indent:
        p.paragraph_format.first_line_indent = Pt(cfg["size_pt"] * para["first_line_indent_chars"])
    run = p.add_run(text)
    _set_run_font(run, cfg["name"], cfg["size_pt"], cfg["bold"])
    return p


def add_separator(doc: Document, fmt: dict):
    """添加分隔线（红色横线，用于版头与正文之间）"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "FF0000")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    return p


def add_page_number(doc: Document, fmt: dict):
    """
    添加页码，格式如 —1—，四号宋体
    单页靠右，双页靠左（通过奇偶页不同页脚实现）
    """
    cfg     = fmt["fonts"]["footer"]
    pn_fmt  = fmt["page_number"]["format"]
    odd_al  = _align_map(fmt["page_number"]["odd_align"])
    even_al = _align_map(fmt["page_number"]["even_align"])

    section = doc.sections[0]
    section.different_first_page_header_footer = False
    section.odd_and_even_pages_header_footer = True

    def _build_footer(footer, align):
        for para in footer.paragraphs:
            para.clear()
        p = footer.paragraphs[0]
        p.alignment = align

        prefix = pn_fmt.split("{n}")[0]
        suffix = pn_fmt.split("{n}")[1] if "{n}" in pn_fmt else ""

        def _run(text):
            r = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            rFonts = OxmlElement("w:rFonts")
            rFonts.set(qn("w:eastAsia"), cfg["name"])
            rFonts.set(qn("w:ascii"), cfg["name"])
            rFonts.set(qn("w:hAnsi"), cfg["name"])
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), str(int(cfg["size_pt"] * 2)))
            rPr.append(rFonts)
            rPr.append(sz)
            t = OxmlElement("w:t")
            t.text = text
            r.append(rPr)
            r.append(t)
            return r

        def _page_field():
            r = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            rFonts = OxmlElement("w:rFonts")
            rFonts.set(qn("w:eastAsia"), cfg["name"])
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), str(int(cfg["size_pt"] * 2)))
            rPr.append(rFonts)
            rPr.append(sz)
            fldChar1 = OxmlElement("w:fldChar")
            fldChar1.set(qn("w:fldCharType"), "begin")
            instr = OxmlElement("w:instrText")
            instr.text = " PAGE "
            fldChar2 = OxmlElement("w:fldChar")
            fldChar2.set(qn("w:fldCharType"), "end")
            r.append(rPr)
            r.append(fldChar1)
            r_instr = OxmlElement("w:r")
            r_instr.append(instr)
            r_end = OxmlElement("w:r")
            r_end.append(fldChar2)
            p._p.append(r)
            p._p.append(r_instr)
            p._p.append(r_end)

        p._p.append(_run(prefix))
        _page_field()
        if suffix:
            p._p.append(_run(suffix))

    _build_footer(section.footer,      odd_al)
    _build_footer(section.even_page_footer, even_al)


def add_image(doc: Document, image_path: str, width_cm: float = 14.0, align: str = "center") -> object:
    """插入图片并居中"""
    from docx.shared import Cm as DocxCm
    p = doc.add_paragraph()
    p.alignment = _align_map(align)
    run = p.add_run()
    run.add_picture(image_path, width=DocxCm(width_cm))
    return p
